from docx import Document
from docx.shared import Pt
from pathlib import Path

output_path = Path(__file__).resolve().parents[1] / "Voice_Recording_Transcript_Consent_Form.docx"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

add_title(doc, "Voice Recording and Transcript Consent Form and Agreement")

doc.add_paragraph("Project Title: [Project Title]")
doc.add_paragraph("Principal Investigator / Lead: [Name, Title]")
doc.add_paragraph("Course / Class: [Machine Learning Class Name and Section]")
doc.add_paragraph("Institution / Organization: [Institution]")
doc.add_paragraph("Contact Email / Phone: [Email / Phone]")
doc.add_paragraph("Protocol / Study ID (if any): [ID]")

add_heading(doc, "1) Purpose of the Project")
doc.add_paragraph(
    "You are being asked to provide voice recordings and transcripts for a research dataset used for "
    "[brief purpose, e.g., speech recognition and language technology research]."
)

add_heading(doc, "2) What You Will Provide")
doc.add_paragraph("By signing this agreement, you confirm that you will provide:")
doc.add_paragraph("- Audio recordings of your own voice; and")
doc.add_paragraph("- Transcripts corresponding to those recordings.")

add_heading(doc, "3) Ownership and Originality")
doc.add_paragraph("You represent and warrant that:")
doc.add_paragraph("- The recordings are of your own voice;")
doc.add_paragraph("- The transcripts are based on your own spoken words; and")
doc.add_paragraph("- You have the right to provide these materials for research use.")

add_heading(doc, "4) Consent to Use and Distribution")
doc.add_paragraph("You grant [Institution/Project Name] a non-exclusive, worldwide, royalty-free license to:")
doc.add_paragraph("- Use, store, copy, and modify the recordings and transcripts for research and educational purposes;")
doc.add_paragraph("- Combine them with other data;")
doc.add_paragraph("- Release and distribute the dataset to other researchers or institutions for non-commercial use.")

add_heading(doc, "5) Non-Commercial Restriction")
doc.add_paragraph("By signing this form, you understand and agree that:")
doc.add_paragraph("- Your contributed audio and transcript data may be released for free non-commercial research and educational use;")
doc.add_paragraph("- Commercial use of your contributed data is not permitted under this repository policy.")

add_heading(doc, "6) Voice Cloning and Voice Synthesis Prohibition")
doc.add_paragraph("You understand and agree that the public is not allowed to use your contributed data to:")
doc.add_paragraph("- Clone, synthesize, or impersonate your voice;")
doc.add_paragraph("- Train or deploy speaker-cloning, voice conversion, or similar identity-mimicking systems for your voice.")

add_heading(doc, "7) Repository Scope and Fork Policy")
doc.add_paragraph("You understand and agree that this main repository is limited to:")
doc.add_paragraph("- Davao Cebuano language data and workflows; and")
doc.add_paragraph("- Whisper model-based ASR workflows.")
doc.add_paragraph("You understand that work on other Philippine languages or non-Whisper model families should be handled in separate forks.")

add_heading(doc, "8) Privacy and Identifiability")
doc.add_paragraph("- Your name and contact details will be kept confidential by the project team.")
doc.add_paragraph("- The dataset may include your voice, which can be personally identifiable.")
doc.add_paragraph("- The dataset may be shared publicly or with other research groups as described above.")

add_heading(doc, "9) Philippine Data Privacy Compliance")
doc.add_paragraph("You acknowledge that collection, storage, processing, and sharing of your data are intended to comply with Philippine law, including Republic Act No. 10173 (Data Privacy Act of 2012), its Implementing Rules and Regulations, and related National Privacy Commission issuances.")

add_heading(doc, "10) Voluntary Participation and Academic Requirement Context")
doc.add_paragraph(
    "Participation is your personal choice and is voluntary. This form confirms your decision to participate "
    "as part of your academic context in this Machine Learning class."
)
doc.add_paragraph(
    "By signing, you acknowledge that you are choosing to forego your identified topic requirement and instead "
    "participate as a research assistant in this class research project."
)
doc.add_paragraph(
    "You further acknowledge that your participation, recordings, and transcripts under this agreement will be "
    "treated as part of your academic requirement for the class, subject to instructor and institutional policies."
)

add_heading(doc, "11) Compensation")
doc.add_paragraph("There is no compensation for participation.")

add_heading(doc, "12) Risks")
doc.add_paragraph("Risks are minimal, but your voice may be recognizable to others if the dataset is shared.")

add_heading(doc, "13) Questions")
doc.add_paragraph("If you have questions about this project, contact:")
doc.add_paragraph("[Name, Email, Phone]")

doc.add_paragraph("")
add_heading(doc, "Participant Statement")
doc.add_paragraph(
    "By signing below, I confirm that I have read and understood this form. I voluntarily consent to provide "
    "recordings and transcripts of my own voice and understand that the resulting dataset may be distributed "
    "for non-commercial research and educational use. I understand that commercial use of my contributed "
    "dataset material is not allowed under repository policy. I understand that public voice cloning, "
    "voice synthesis, or voice impersonation of my contributed voice is prohibited. I understand that this "
    "main repository is scoped to Davao Cebuano and Whisper-based workflows, and that other languages or model families should be "
    "handled in separate forks. I acknowledge the project's intent to comply with Philippine data privacy law, "
    "including Republic Act No. 10173. I understand there is no monetary compensation for participation. I also "
    "understand and agree that this participation is in lieu of my identified topic and is part of my academic "
    "requirement in the Machine Learning class, as stated above."
)

doc.add_paragraph("")
doc.add_paragraph("Participant Name: ___________________________")
doc.add_paragraph("Program / Year / Section: ____________________")
doc.add_paragraph("Signature: _________________________________")
doc.add_paragraph("Date: ___________________")

doc.add_paragraph("")
doc.add_paragraph("Instructor / Research Lead Name: __________________")
doc.add_paragraph("Signature: ______________________________________")
doc.add_paragraph("Date: ___________________")

doc.save(output_path)
print(f"Updated: {output_path}")
